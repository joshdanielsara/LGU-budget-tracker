from django.shortcuts import render, redirect, get_object_or_404  # Render templates, redirect views, retrieve an object or 404 if not found
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse  # HTTP response types for returning plain text, redirection, JSON data
from django.db.models import F, ExpressionWrapper, FloatField  # Allows field expressions and calculations directly in database queries
from django.contrib.auth.decorators import login_required  # Decorator to restrict view access to authenticated users only
from django.core.paginator import Paginator  # Paginator for splitting large datasets across multiple pages
from django.core.cache import cache  # Cache framework to store and retrieve data for performance optimization
from io import BytesIO  # In-memory byte streams, useful for temporary file handling
from docx import Document  # For creating and manipulating Word documents
from .firebase import database  # Imports Firebase database helper functions from local `firebase` module
import firebase_admin  # Firebase Admin SDK for managing Firebase services
from firebase_admin import credentials, db  # For Firebase credentials and accessing the Firebase database
import datetime  # Module for handling date and time operations
import re  # Regular expressions for pattern matching in strings
from easyaudit.models import CRUDEvent  # Model for logging Create, Read, Update, Delete (CRUD) events
from django.contrib.contenttypes.models import ContentType  # Allows referencing Django models generically using content types
from django.contrib.auth.models import User  # Django's built-in User model for user management
from KwentasApp.models import FirebaseEntry  # Custom model representing an entry in Firebase database
from django.contrib import messages
from time import time
from datetime import datetime

def create_entry(request):
    if request.method == 'POST':
        # Extract data from the form submission
        ppa = request.POST.get('ppa')
        implementing_unit = request.POST.get('implementing_unit')
        location = request.POST.get('location')        
        appropriation_str = request.POST.get('appropriation', '0').replace(',', '') 
        remarks = request.POST.get('remarks')
        start_date_str = request.POST.get('start_date')
        end_date_str = request.POST.get('end_date')
        code = request.POST.get('code')
        services = request.POST.get('services')
        year_str = request.POST.get('year')

        # Validate required fields
        if not (ppa and implementing_unit and location and appropriation_str and start_date_str and end_date_str and code and services and year_str):
            messages.error(request, "All fields are required.")
            return redirect('create_entry')

        # Validate appropriation amount
        try:
            appropriation = float(appropriation_str)
            if appropriation <= 0:
                messages.error(request, "Appropriation amount must be a positive number.")
                return redirect('create_entry')
        except ValueError:
            messages.error(request, "Invalid appropriation amount.")
            return redirect('create_entry')

        # Validate date format
        date_format = r'\d{4}-\d{2}-\d{2}'  # YYYY-MM-DD
        if not (re.match(date_format, start_date_str) and re.match(date_format, end_date_str)):
            messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
            return redirect('create_entry')

        # Convert dates to datetime objects for further validation
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d')
        except ValueError:
            messages.error(request, "Invalid date.")
            return redirect('create_entry')

        # Validate start date is before end date
        if start_date >= end_date:
            messages.error(request, "Start date must be before end date.")
            return redirect('create_entry')

        # Validate code format
        if not re.match(r'^[\w\d\s-]+$', code):
            messages.error(request, "Invalid code format. Only letters, numbers, spaces, and hyphens are allowed.")
            return redirect('create_entry')

        # Validate code uniqueness
        if database.child('Data').child(code).get().val() is not None:
            messages.error(request, "Code already exists.")
            return redirect('create_entry')

        # Validate year format
        if not re.match(r'\d{4}', year_str):
            messages.error(request, "Invalid year format. Please use YYYY.")
            return redirect('create_entry')

        # Set remaining balance equal to appropriation
        overall_budget = appropriation
        remaining_total_balance = overall_budget
        total_disbursements = 0
        total_obligations = 0
        remaining_obligations = 0

        # Calculate the utilization rate (initially 0 since total_disbursements is 0)
        utilization_rate = (total_disbursements / overall_budget) * 100 if overall_budget > 0 else 0

        try:

            created_at = time()
            # Save project entry with code as the key
            new_entry_ref = database.child('Data').child(code).set({
                "ppa": ppa,
                "implementing_unit": implementing_unit,
                "location": location,
                "appropriation": appropriation,
                "remarks": remarks,
                "start_date": start_date_str,
                "end_date": end_date_str,
                "code": code,
                "services": services,
                "year": year_str,
                "remaining_total_balance": remaining_total_balance,
                "total_disbursements": total_disbursements,  # Set total spent default value to 0
                "added_budget": 0,  # Set added budget default value to 0
                "overall_budget": overall_budget,
                "utilization_rate": utilization_rate,  # Include utilization rate
                "total_obligations": total_obligations,
                "remaining_obligations": remaining_obligations,
                 "created_at": created_at,
            })

            CRUDEvent.objects.create(
                event_type=CRUDEvent.CREATE,
                object_id=code,  # Reference the Firebase "code" as object_id
                object_repr=f"Project Entry: {code} is added",  # Description of the object
                content_type=ContentType.objects.get(app_label='KwentasApp', model='firebaseentry'),  # Use FirebaseEntry model for logging
                user=request.user if request.user.is_authenticated else None,
            )

            # Refresh the cache with the new entry
            cache.delete('project_entries')
            get_project_entries()

            messages.success(request, "Successfully added")
            return redirect('create_entry')

        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect('create_entry')

    else:
        return render(request, 'KwentasApp/adddata.html')





def get_project_entries():
    # Check if data is cached
    cached_data = cache.get('project_entries')
    if cached_data:
        return cached_data

    result = database.child('Data').get()

    entries_below_2024 = []
    entries_2024_and_above = []
    all_entries = []

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue  # Skip the placeholder entry

            remarks = value.get('remarks')  # Extract remarks
            is_awarded = remarks == "Awarded Already"  # Determine if awarded

            # Parse created_at as an integer or default to 0 if it's missing or invalid
            created_at = value.get('created_at')
            try:
                created_at = int(created_at) if created_at is not None else 0
            except ValueError:
                created_at = 0  # Handle cases where created_at isn't a valid integer

            entry = {
                'ppa': value.get('ppa'),
                'implementing_unit': value.get('implementing_unit'),
                'location': value.get('location'),
                'appropriation': value.get('appropriation'),
                'remarks': remarks,
                'start_date': value.get('start_date'),
                'end_date': value.get('end_date'),
                'code': value.get('code'),
                'services': value.get('services'),
                'year': value.get('year'),
                'remaining_total_balance': value.get('remaining_total_balance'),
                'total_disbursements': value.get('total_disbursements'),
                'added_budget': value.get('added_budget'),
                'overall_budget': value.get('overall_budget'),
                'utilization_rate': value.get('utilization_rate'),
                'total_obligations': value.get('total_obligations'),
                'remaining_obligations': value.get('remaining_obligations'),
                'created_at': created_at,  # Store converted created_at
                'disbursement': [],
                'obligation': [],  # Initialize obligation list
                'budget_data': [],  # Initialize budget data list
                'is_awarded': is_awarded  # Include awarded status
            }

            # Check if obligation node exists
            if 'obligation' in value:
                for obligation_key, obligation_value in value['obligation'].items():
                    entry['obligation'].append({
                        'name': obligation_value.get('name'),
                        'obligation': obligation_value.get('obligation'),
                        'date': obligation_value.get('date')
                    })

            # Check if disbursement node exists
            if 'disbursement' in value:
                for disbursement_key, disbursement_value in value['disbursement'].items():
                    entry['disbursement'].append({
                        'key': disbursement_key,  # Add the key here
                        'name': disbursement_value.get('name'),
                        'voucher': disbursement_value.get('voucher'),
                        'disbursement': disbursement_value.get('disbursement'),
                        'date': disbursement_value.get('date'),
                        'status': disbursement_value.get('status'),
                        'check_date': disbursement_value.get('check_date')
                    })

            # Fetch budget data for the current entry
            budget_data = database.child('Data').child(key).child('budget').get().val()
            if budget_data:
                for budget_key, budget_value in budget_data.items():
                    entry['budget_data'].append({
                        'name': budget_value.get('name'),
                        'added_budget': budget_value.get('added_budget'),
                        'date': budget_value.get('date')
                    })

            all_entries.append(entry)

            # Check the year and add to the appropriate list
            if entry.get('year') is not None:
                if int(entry['year']) < 2024:
                    entries_below_2024.append(entry)
                else:
                    entries_2024_and_above.append(entry)

    # Sort entries by 'created_at' in descending order
    entries_below_2024.sort(key=lambda x: x['created_at'], reverse=True)
    entries_2024_and_above.sort(key=lambda x: x['created_at'], reverse=True)
    all_entries.sort(key=lambda x: x['created_at'], reverse=True)

    # Cache the data for future use (e.g., 1 hour)
    cache.set('project_entries', (entries_below_2024, entries_2024_and_above, all_entries), 3600)

    return entries_below_2024, entries_2024_and_above, all_entries



def continuing_projects(request):
    # Retrieve project entries for continuing projects
    entries_below_2024, _, all_entries = get_project_entries()
    
    # Get sorting option from query parameters; do nothing if not specified
    sort_option = request.GET.get('sort')
    
    # Apply sorting if an option is selected
    if sort_option == 'oldest':
        entries_below_2024.sort(key=lambda x: x['created_at'])  # Sort by oldest
    elif sort_option == 'recent':
        entries_below_2024.sort(key=lambda x: x['created_at'], reverse=True)  # Sort by recent
    
    # Paginate the entries (sorted or unsorted)
    paginator = Paginator(entries_below_2024, 10)  # Show 10 projects per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Render the page with sorted (or unsorted) project entries
    return render(request, 'KwentasApp/continuing.html', {
        'page_obj': page_obj,
        'all_entries': all_entries,
        'sort_option': sort_option  # Pass current sort option to the template
    })



def current_projects(request):
    # Retrieve project entries
    entries_below_2024, entries_2024_and_above, all_entries = get_project_entries()
    
    # Get sorting option from query parameters; do nothing if not specified
    sort_option = request.GET.get('sort')
    
    # Apply sorting if an option is selected
    if sort_option == 'oldest':
        entries_2024_and_above.sort(key=lambda x: x['created_at'])  # Sort by oldest
    elif sort_option == 'recent':
        entries_2024_and_above.sort(key=lambda x: x['created_at'], reverse=True)  # Sort by recent
    
    # Paginate the entries (sorted or unsorted)
    paginator = Paginator(entries_2024_and_above, 10)  # Show 10 projects per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'KwentasApp/current.html', {
        'page_obj': page_obj,
        'all_entries': all_entries,
        'sort_option': sort_option  # Pass current sort option to the template
    })
from django.core.paginator import Paginator
from django.shortcuts import render

def handle_entries(request, template_name):
    """
    Handles retrieving, sorting, and paginating project entries
    and renders the specified template.
    """
    # Retrieve project entries
    _, _, all_entries = get_project_entries()
    
    # Get sorting option from query parameters; do nothing if not specified
    sort_option = request.GET.get('sort')
    
    # Apply sorting if an option is selected
    if sort_option == 'oldest':
        all_entries.sort(key=lambda x: x['created_at'])  # Sort by oldest
    elif sort_option == 'recent':
        all_entries.sort(key=lambda x: x['created_at'], reverse=True)  # Sort by recent
    
    # Paginate the entries (sorted or unsorted)
    paginator = Paginator(all_entries, 10)  # Show 10 projects per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Render the appropriate template
    return render(request, template_name, {
        'page_obj': page_obj,
        'all_entries': all_entries,
        'sort_option': sort_option  # Pass current sort option to the template
    })


def obligations(request):
    return handle_entries(request, 'KwentasApp/obligations.html')

def disbursements(request):
    return handle_entries(request, 'KwentasApp/disbursements.html')

def procurements(request):
    return handle_entries(request, 'KwentasApp/procurements.html')

def check_payment(request):
    return handle_entries(request, 'KwentasApp/check_payment.html')












def add_obligation(request, project_type):
    if request.method == 'POST':
        entry_key = request.POST.get('entry-code')
        name = request.POST.get('obligation_name')
        obligation_input = request.POST.get('obligation_input', '0')
        date = request.POST.get('obligation_date')

        # Determine redirect URL based on project type
        if project_type == 'continuing':
            redirect_url = '/continuing_projects'
        elif project_type == 'current':
            redirect_url = '/current_projects'
        else:
            redirect_url = '/'

        try:
            # Remove any commas or spaces from the input
            spent_cleaned = re.sub(r'[,\s]', '', obligation_input)

            # Validate that the cleaned input is a valid number
            if not re.match(r'^\d+(\.\d{1,2})?$', spent_cleaned):
                messages.error(request, "Invalid 'spent' value. It should only contain numbers, without commas or spaces.")
                return redirect(redirect_url)

            # Convert the cleaned and validated string to a float
            obligation = float(spent_cleaned)

            entry_ref = database.child('Data').child(entry_key)  # Assigning database child to entry_ref

            # Get data for the entry
            entry_data = entry_ref.get().val()

            # Check if entry_data is None
            if entry_data is None:
                messages.error(request, "Entry not found.")
                return redirect(redirect_url)

            start_date = entry_data.get('start_date')
            end_date = entry_data.get('end_date')

            # Ensure date, start_date, and end_date are properly formatted as date objects
            try:
                date = datetime.strptime(date, '%Y-%m-%d').date()
                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect(redirect_url)

            # Validate date range
            if date < start_date:
                messages.error(request, "Date must not be before the project's start date.")
                return redirect(redirect_url)

            if date > end_date:
                messages.error(request, "Date must not be after the project's end date.")
                return redirect(redirect_url)

            # Update total spent for the entry
            total_obligations = entry_data.get('total_obligations', 0) + obligation
            total_disbursements = entry_data.get('total_disbursements')

            # Calculate remaining balance
            overall_budget = entry_data.get('overall_budget', 0)
            remaining_total_balance = overall_budget - total_obligations
            remaining_obligations = total_obligations - total_disbursements

            if remaining_total_balance < 0:
                messages.error(request, "Remaining Balance cannot be negative. Add Budget if you wish to continue.")
                return redirect(redirect_url)

            # Push a new child node with a unique key for the obligation under the chosen entry
            new_obligation_ref = database.child('Data').child(entry_key).child('obligation').push({
                'name': name,
                'obligation': obligation,
                'date': date.strftime('%Y-%m-%d')  # Store the date as a string
            })

            # Update total spent and remaining balance under the entry
            database.child('Data').child(entry_key).update({
                "total_obligations": total_obligations,
                "remaining_obligations": remaining_obligations,
                 "remaining_total_balance": remaining_total_balance
            })

            CRUDEvent.objects.create(
                event_type=CRUDEvent.UPDATE,
                object_id=entry_key,  # Reference the Firebase "code" as object_id
                object_repr=f"Project Entry: {entry_key} - Obligation Name: {name} with Amount of: {obligation}",  # Include both name and amount in the log
                content_type=ContentType.objects.get(app_label='KwentasApp', model='firebaseentry'),  # Replace with your actual model
                user=request.user if request.user.is_authenticated else None
            )

            # Invalidate the cache
            cache.delete('project_entries')

            messages.success(request, "Obligation added successfully.")
            return redirect(redirect_url)
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect(redirect_url)
    else:
        messages.error(request, "Method not allowed")
        return redirect(redirect_url)

def add_disbursement(request, project_type):
    if request.method == 'POST':
        entry_key = request.POST.get('entry-code')
        voucher = request.POST.get('voucher')
        name = request.POST.get('disbursement_name')
        disbursement_input = request.POST.get('disbursement_input', '0')
        date = request.POST.get('disbursement_date')

        # Determine redirect URL based on project type
        if project_type == 'disbursement':
            redirect_url = '/disbursements'
        else:
            redirect_url = '/'

        try:
            # Remove any commas or spaces from the input
            spent_cleaned = re.sub(r'[,\s]', '', disbursement_input)

            # Validate that the cleaned input is a valid number
            if not re.match(r'^\d+(\.\d{1,2})?$', spent_cleaned):
                messages.error(request, "Invalid 'spent' value. It should only contain numbers, without commas or spaces.")
                return redirect(redirect_url)

            # Convert the cleaned and validated string to a float
            disbursement = float(spent_cleaned)
            
            entry_ref = database.child('Data').child(entry_key)  # Assigning database child to entry_ref

            # Get data for the entry
            entry_data = entry_ref.get().val()

            # Check if entry_data is None
            if entry_data is None:
                messages.error(request, "Entry not found.")
                return redirect(redirect_url)

            start_date = entry_data.get('start_date')
            end_date = entry_data.get('end_date')

            # Ensure date, start_date, and end_date are properly formatted as date objects
            try:
                date = datetime.strptime(date, '%Y-%m-%d').date()
                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect(redirect_url)

            # Validate date range
            if date < start_date:
                messages.error(request, "Date must not be before the project's start date.")
                return redirect(redirect_url)

            if date > end_date:
                messages.error(request, "Date must not be after the project's end date.")
                return redirect(redirect_url)

            # Update total spent for the entry
            remaining_balance = entry_data.get('remaining_total_balance')
            total_disbursements = entry_data.get('total_disbursements', 0) + disbursement
            total_obligations = entry_data.get('total_obligations')
            
            # Calculate remaining balance of obligation
            overall_budget = entry_data.get('overall_budget', 0)
            remaining_obligations = total_obligations - total_disbursements
           
            if remaining_obligations < 0:
                messages.error(request, "Remaining Obligations cannot be negative. Add Obligation if you wish to continue.")
                return redirect(redirect_url)

            # Push a new child node with a unique key for the disbursement under the chosen entry
          # Push a new child node with a unique key for the disbursement under the chosen entry
            new_disbursement_ref = database.child('Data').child(entry_key).child('disbursement').push({
                'name': name,
                'voucher': voucher,
                'disbursement': disbursement,
                'date': date.strftime('%Y-%m-%d'),  # Store the date as a string
                'status': 'Pending',  # Set default status to "Pending"
                'check_date': date.strftime('%Y-%m-%d')  # Set check payment date same as the disbursement date
            })


            # Update total disbursements and remaining balance under the entry
            database.child('Data').child(entry_key).update({
                "total_disbursements": total_disbursements,
                "remaining_obligations": remaining_obligations
            })

            # Calculate and update utilization rate
            utilization_rate = (total_disbursements / overall_budget) * 100 if overall_budget > 0 else 0
            database.child('Data').child(entry_key).update({"utilization_rate": utilization_rate})

            # Log the CRUD event
            content_type = ContentType.objects.get(app_label='KwentasApp', model='firebaseentry')  # Replace 'firebaseentry' with the actual model name
            CRUDEvent.objects.create(
                event_type=CRUDEvent.UPDATE,
                object_id=entry_key,  # Reference the Firebase "code" as object_id
                object_repr=f"Project Entry: {entry_key} - Disbursement Name: {name} (with Amount of: {disbursement})",  # Include both name and amount in the log
                content_type=content_type,
                user=request.user if request.user.is_authenticated else None
            )

            # Invalidate the cache
            cache.delete('project_entries')

            # Display success message and redirect
            messages.success(request, "Disbursement added successfully.")
            return redirect(redirect_url)

        except Exception as e:
            # Handle unexpected errors
            messages.error(request, f"Error: {str(e)}")
            return redirect(redirect_url)

    else:
        # Method not allowed
        messages.error(request, "Method not allowed.")
        return redirect(redirect_url)



def update_disbursement(request):
    if request.method == 'POST':
        entry_code = request.POST.get('entry-code')
        disbursement_key = request.POST.get('disbursement-key')
        new_status = request.POST.get('status')
        new_check_date = request.POST.get('check_date')

        # Validate inputs
        if not entry_code or not disbursement_key:
            messages.error(request, "Missing entry or disbursement identifier.")
            return redirect(request.META.get('HTTP_REFERER', '/'))

        try:
            # Fetch current disbursement data
            current_disbursement = database.child('Data').child(entry_code).child('disbursement').child(disbursement_key).get().val()
            if not current_disbursement:
                messages.error(request, "Disbursement not found.")
                return redirect(request.META.get('HTTP_REFERER', '/'))

            # Update fields
            updated_data = {
                'status': new_status or current_disbursement.get('status', 'Pending'),
                'check_date': new_check_date or current_disbursement.get('check_date', None)
            }

            # Update Firebase
            database.child('Data').child(entry_code).child('disbursement').child(disbursement_key).update(updated_data)

            # Clear cache after update
            cache.delete('project_entries')  # Clear the cached project entries

            # Success message
            messages.success(request, "Disbursement updated successfully.")
            return redirect(request.META.get('HTTP_REFERER', '/'))

        except Exception as e:
            messages.error(request, f"Error updating disbursement: {str(e)}")
            return redirect(request.META.get('HTTP_REFERER', '/'))

    messages.error(request, "Invalid request method.")
    return redirect('/')

 

    

def add_budget(request, project_type):
    if request.method == 'POST':
        entry_key = request.POST.get('entry_code')
        name = request.POST.get('budget_name')
        added_budget_str = request.POST.get('added_budget', '0')  # Fetching it as a string for validation
        date = request.POST.get('budget_date')

        # Determine redirect URL based on project type
        if project_type == 'continuing':
            redirect_url = '/continuing_projects'
        elif project_type == 'current':
            redirect_url = '/current_projects'
        else:
            redirect_url = '/'

        try:
            # Remove commas or spaces from added_budget input
            added_budget_str = re.sub(r'[,\s]', '', added_budget_str)

            # Validate that the cleaned input is a valid number
            if not re.match(r'^\d+(\.\d{1,2})?$', added_budget_str):
                messages.error(request, "Invalid 'added budget' value. It should only contain numbers, without commas or spaces.")
                return redirect(redirect_url)

            # Convert the cleaned and validated string to a float
            added_budget = float(added_budget_str)

            # Get data for the entry
            entry_ref = database.child('Data').child(entry_key)  # Assigning database child to entry_ref
            entry_data = entry_ref.get().val()

            # Check if entry_data is None
            if entry_data is None:
                messages.error(request, "Entry not found.")
                return redirect(redirect_url)

            start_date = entry_data.get('start_date')
            end_date = entry_data.get('end_date')

            # Ensure date, start_date, and end_date are properly formatted as date objects
            try:
                date = datetime.strptime(date, '%Y-%m-%d').date()
                start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect(redirect_url)

            # Validate date range
            if date < start_date:
                messages.error(request, "Date must not be before the project's start date.")
                return redirect(redirect_url)

            if date > end_date:
                messages.error(request, "Date must not be after the project's end date.")
                return redirect(redirect_url)

            # Update total added budget for the entry
            total_added_budget = entry_data.get('added_budget', 0) + added_budget
            overall_budget = entry_data.get('overall_budget', 0)
            overall_budget += total_added_budget
            remaining_total_balance = overall_budget - total_added_budget
            

            # Check if the remaining balance would go negative
            if remaining_total_balance < 0:
                messages.error(request, "Remaining Balance cannot be negative. Add Budget if you wish to continue.")
                return redirect(redirect_url)

            # Push a new child node with a unique key for the added budget under the chosen entry
            new_budget_ref = database.child('Data').child(entry_key).child('budget').push({
                'name': name,
                'added_budget': added_budget,
                'date': date.strftime('%Y-%m-%d')  # Store the date as a string
            })

            # Update total added budget under the entry
            database.child('Data').child(entry_key).update({
                "added_budget": total_added_budget,
                "remaining_total_balance": remaining_total_balance,
                "overall_budget": overall_budget,
            })

            # Calculate utilization rate
            total_disbursements = entry_data.get('total_disbursements', 0)
            utilization_rate = (total_disbursements / overall_budget) * 100 if overall_budget > 0 else 0
            database.child('Data').child(entry_key).update({"utilization_rate": utilization_rate})

            # Log the event
            content_type = ContentType.objects.get(app_label='KwentasApp', model='firebaseentry')  # Adjust model name
            CRUDEvent.objects.create(
                event_type=CRUDEvent.UPDATE,
                object_id=entry_key,  # Reference the Firebase "code" as object_id
                object_repr=f"Project Entry: {entry_key} - Budget Added: {name} (with Amount of: {added_budget})",
                content_type=content_type,
                user=request.user if request.user.is_authenticated else None
            )

            # Invalidate the cache
            cache.delete('project_entries')

            messages.success(request, "Budget added successfully.")
            return redirect(redirect_url)
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect(redirect_url)
    else:
        messages.error(request, "Method not allowed")
        return redirect(redirect_url)


def search_projects(request, project_type):
    query = request.GET.get('query', '').lower()
    entries = []
    matched_entries = []

    # Determine the appropriate list and redirect URL based on project type
    if project_type == 'continuing':
        entries, _, all_entries = get_project_entries()
        year_condition = lambda year: year is not None and int(year) < 2024
        template = 'KwentasApp/continuing.html'
    elif project_type == 'current':
        _, entries, all_entries = get_project_entries()
        year_condition = lambda year: year is not None and int(year) >= 2024
        template = 'KwentasApp/current.html'
    elif project_type == 'obligations':
        _, entries, all_entries = get_project_entries()
        year_condition = lambda year: year is not None
        template = 'KwentasApp/obligations.html'
    elif project_type == 'disbursements':
        _, entries, all_entries = get_project_entries()
        year_condition = lambda year: year is not None
        template = 'KwentasApp/disbursements.html'
    elif project_type == 'procurements':
        _, entries, all_entries = get_project_entries()
        year_condition = lambda year: year is not None
        template = 'KwentasApp/procurements.html'
    else:
        return JsonResponse({'error': 'Invalid project type'}, status=400)

    # Fetch data from Firebase
    result = database.child('Data').get()

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue  # Skip the placeholder entry

            remarks = value.get('remarks')  # Extract remarks
            is_awarded = remarks == "Awarded Already"  # Determine if awarded
            
            entry = {
                'ppa': value.get('ppa'),
                'implementing_unit': value.get('implementing_unit'),
                'location': value.get('location'),
                'appropriation': value.get('appropriation'),
                'remarks': value.get('remarks'),
                'start_date': value.get('start_date'),
                'end_date': value.get('end_date'),
                'code': value.get('code'),
                'services': value.get('services'),
                'year': value.get('year'),
                'remaining_total_balance': value.get('remaining_total_balance'),
                'total_disbursements': value.get('total_disbursements'),
                'added_budget': value.get('added_budget'),
                'overall_budget': value.get('overall_budget'),
                'utilization_rate': value.get('utilization_rate'),
                'total_obligations': value.get('total_obligations'),
                'remaining_obligations': value.get('remaining_obligations'),
                'is_awarded': is_awarded,
                'disbursement': [],
                'obligation': [],  # Initialize obligation list
                'budget_data': [],  # Initialize budget data list
            }

            # Search and filter based on project type
            if any(query in str(entry[field]).lower() for field in [
                'ppa', 'implementing_unit', 'location', 'appropriation', 'remarks',
                'start_date', 'end_date', 'code', 'services'
            ]):
                if year_condition(entry.get('year')):
                    matched_entries.append(entry)

    # Return JSON response if AJAX request
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'results': matched_entries})

    # Render the template for normal requests
    return render(request, template, {
        'matched_entries': matched_entries,
        'query': query,
        'entries': entries,
        'all_entries': all_entries
    })


def update_entry(request, project_type):
    if request.method == 'POST':
        entry_key = request.POST.get('entry_code')
        ppa = request.POST.get('ppa', '')  
        implementing_unit = request.POST.get('implementing_unit', '')  
        start_date_str = request.POST.get('start_date', '') 
        end_date_str = request.POST.get('end_date', '') 
        year_str = request.POST.get('year', '') 
        code = request.POST.get('code', '')  # Added code field
        services = request.POST.get('services', '')  # Added services field
        remarks = request.POST.get('remarks', '')  # Added remarks field
        location = request.POST.get('location', '')  # Added location field

        # Determine redirect URL based on project type
        if project_type == 'current':
            redirect_url = '/current_projects'
        elif project_type == 'continuing':
            redirect_url = '/continuing_projects'
        else:
            redirect_url = '/'

        # Validate date format if not empty
        if start_date_str and end_date_str:
            date_format = r'\d{4}-\d{2}-\d{2}'  # YYYY-MM-DD
            if not (re.match(date_format, start_date_str) and re.match(date_format, end_date_str)):
                messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
                return redirect(redirect_url)

            try:
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                end_date = datetime.strptime(end_date_str, '%Y-%m-%d')
                if start_date >= end_date:
                    messages.error(request, "Start date must be before end date.")
                    return redirect(redirect_url)
            except ValueError:
                messages.error(request, "Invalid date.")
                return redirect(redirect_url)

        # Validate year format if not empty
        if year_str and not re.match(r'\d{4}', year_str):
            messages.error(request, "Invalid year format. Please use YYYY.")
            return redirect(redirect_url)

        try:
            # Get data for the entry
            entry_data = database.child('Data').child(entry_key).get().val()
            if entry_data is None:
                messages.error(request, "Entry not found.")
                return redirect(redirect_url)

            # Check if the code is being changed and if it already exists
            if code != entry_key and database.child('Data').child(code).get().val() is not None:
                messages.error(request, "Code already exists.")
                return redirect(redirect_url)

            updated_fields = []
            
            # Track changes
            if ppa != entry_data.get('ppa', ''):
                updated_fields.append(f"PPA: {entry_data.get('ppa', '')} → {ppa}")
            if implementing_unit != entry_data.get('implementing_unit', ''):
                updated_fields.append(f"Implementing Unit: {entry_data.get('implementing_unit', '')} → {implementing_unit}")
            if start_date_str != entry_data.get('start_date', ''):
                updated_fields.append(f"Start Date: {entry_data.get('start_date', '')} → {start_date_str}")
            if end_date_str != entry_data.get('end_date', ''):
                updated_fields.append(f"End Date: {entry_data.get('end_date', '')} → {end_date_str}")
            if year_str != entry_data.get('year', ''):
                updated_fields.append(f"Year: {entry_data.get('year', '')} → {year_str}")
            if code != entry_key:
                updated_fields.append(f"Code: {entry_key} → {code}")
            if services != entry_data.get('services', ''):
                updated_fields.append(f"Services: {entry_data.get('services', '')} → {services}")
            if location != entry_data.get('location', ''):
                updated_fields.append(f"Location: {entry_data.get('location', '')} → {location}")
            if remarks != entry_data.get('remarks', ''):
                updated_fields.append(f"Remarks: {entry_data.get('remarks', '')} → {remarks}")

            updated_fields_str = ', '.join(updated_fields) if updated_fields else 'No fields updated'

            # Handle code update case
            if code != entry_key:
                new_entry_data = {
                    "ppa": ppa,
                    "implementing_unit": implementing_unit,
                    "start_date": start_date_str,
                    "end_date": end_date_str,
                    "year": year_str,
                    "code": code,
                    "services": services,
                    "location": location,
                    "remarks": remarks,
                    "remaining_total_balance": entry_data.get('remaining_total_balance', 0),
                    "total_disbursements": entry_data.get('total_disbursements', 0),
                    "added_budget": entry_data.get('added_budget', 0),
                    "overall_budget": entry_data.get('overall_budget', 0),
                    "utilization_rate": entry_data.get('utilization_rate', 0),
                    "total_obligations": entry_data.get('total_obligations', 0),
                    "remaining_obligations": entry_data.get('remaining_obligations', 0),
                }
                database.child('Data').child(code).set(new_entry_data)
                database.child('Data').child(entry_key).remove()
            else:
                database.child('Data').child(entry_key).update({
                    "ppa": ppa,
                    "implementing_unit": implementing_unit,
                    "start_date": start_date_str,
                    "end_date": end_date_str,
                    "year": year_str,
                    "code": code,
                    "services": services,
                    "location": location,
                    "remarks": remarks
                })

            # Log the update event
            content_type = ContentType.objects.get(app_label='KwentasApp', model='firebaseentry')
            CRUDEvent.objects.create(
                event_type=CRUDEvent.UPDATE,
                object_id=entry_key,
                object_repr=f"Project Entry: {entry_key} - Updated Fields: {updated_fields_str}",
                content_type=content_type,
                user=request.user if request.user.is_authenticated else None
            )

            # Clear cache and add success message
            cache.delete('project_entries')
            messages.success(request, f"Project Entry: {entry_key} has been successfully updated.")

            return redirect(request.META.get('HTTP_REFERER', '/'))
        except Exception as e:
            messages.error(request, f"Error: {str(e)}")
            return redirect(redirect_url)
    else:
        messages.error(request, "Method not allowed.")
        return redirect(redirect_url)


# Function to ensure the placeholder exists
def ensure_placeholder():
    data = database.child('Data').get().val()
    if 'placeholder' not in data:
        database.child('Data').child('placeholder').set(True)

# Add a placeholder initially
def add_placeholder():
    database.child('Data').child('placeholder').set(True)

# Ensure the placeholder exists when the module is loaded
add_placeholder()





def delete_entry(request, project_type):
    if request.method == 'POST':
        entry_key = request.POST.get('entry_code')


        if project_type == 'continuing':
            redirect_url = '/continuing_projects'
        elif project_type == 'current':
            redirect_url = '/current_projects'
        else:
            redirect_url = '/'
        
        try:
            # Delete the specific entry
            database.child('Data').child(entry_key).remove()
            # Ensure placeholder remains
            ensure_placeholder()

            # Refresh the cache after deletion
            cache.delete('project_entries')

            # Log the deletion event
            content_type = ContentType.objects.get(app_label='KwentasApp', model='firebaseentry')  # Adjust model name if necessary
            CRUDEvent.objects.create(
                event_type=CRUDEvent.DELETE,
                object_id=entry_key,  # Reference the Firebase "code" as object_id
                object_repr=f"Project Entry: {entry_key} is Deleted",
                content_type=content_type,
                user=request.user if request.user.is_authenticated else None
            )

            # Add success message
            messages.success(request, f"Project Entry: {entry_key} has been successfully deleted.")
            return redirect(request.META.get('HTTP_REFERER', redirect_url))
            
        except Exception as e:
            # Add error message
            messages.error(request, f"Error: {str(e)}")
            return redirect(redirect_url)
    else:
        return redirect(redirect_url)



from django.contrib.auth.decorators import login_required
@login_required
def reports_view(request):
    _, _, all_entries = get_project_entries()

    total_utilization = 0
    total_entries = 0

    below_50_utilization = []
    above_50_utilization = []

    count_below_50 = 0  # Count for projects below 50% utilization
    count_above_50 = 0  # Count for projects above 50% utilization

    for entry in all_entries:
        if 'utilization_rate' in entry and entry['utilization_rate'] is not None:
            utilization_rate = entry['utilization_rate']
            total_utilization += utilization_rate
            total_entries += 1
            
            # Categorize projects based on utilization rate
            if utilization_rate < 50:
                below_50_utilization.append(entry)
                count_below_50 += 1  # Increment the below 50 count
            else:
                above_50_utilization.append(entry)
                count_above_50 += 1  # Increment the above 50 count

    average_utilization = total_utilization / total_entries if total_entries > 0 else 0

    # Debugging statement to ensure average_utilization and counts are calculated
    print(f"Average Utilization: {average_utilization}")
    print(f"Projects Below 50%: {count_below_50}")
    print(f"Projects Above 50%: {count_above_50}")

    return render(request, 'KwentasApp/stats.html', {
        'average_utilization': average_utilization,
        'below_50_utilization': below_50_utilization,
        'above_50_utilization': above_50_utilization,
        'count_below_50': count_below_50,
        'count_above_50': count_above_50,
    })

@login_required
def graphs(request):
    _, _, all_entries = get_project_entries()

    utilization_rates = []  # List to hold utilization rates
    total_utilization = 0
    total_entries = 0

    for entry in all_entries:
        if 'utilization_rate' in entry and entry['utilization_rate'] is not None:
            utilization_rate = entry['utilization_rate']
            utilization_rates.append(utilization_rate)  # Collect utilization rates
            total_utilization += utilization_rate  # Sum up utilization rates
            total_entries += 1  # Count entries

    # Calculate average utilization rate
    average_utilization = total_utilization / total_entries if total_entries > 0 else 0


    return render(request, 'KwentasApp/graphs.html', {
        'utilization_rates': utilization_rates,
        'average_utilization': average_utilization,  # Pass average utilization to the template
    })

def all_projects(request):
    _, _, all_entries = get_project_entries()
    query = request.GET.get('query', '')  # Default to empty string if query is None
    paginator = Paginator(all_entries, 10)  # Show 10 projects per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    result = database.child('Data').get()

    matched_entries = []

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue  # Skip the placeholder entry
            entry = {
                'ppa': value.get('ppa'),
                'implementing_unit': value.get('implementing_unit'),
                'location': value.get('location'),
                'appropriation': value.get('appropriation'),
                'remarks': value.get('remarks'),
                'start_date': value.get('start_date'),
                'end_date': value.get('end_date'),
                'code': value.get('code'),
                'services': value.get('services'),
                'year': value.get('year'),
                'remaining_total_balance': value.get('remaining_total_balance'),
                'total_disbursements': value.get('total_disbursements'),
                'total_obligations': value.get('total_obligations'),
                'remaining_obligations': value.get('remaining_obligations'),
                'obligation': []  # Initialize obligation list
               
            }

        
            # Perform a case-insensitive search only if the fields are not None
            if (query.lower() in str(entry['ppa']).lower() if entry['ppa'] else False or
                query.lower() in str(entry['implementing_unit']).lower() if entry['implementing_unit'] else False or
                query.lower() in str(entry['location']).lower() if entry['location'] else False or
                query.lower() in str(entry['appropriation']).lower() if entry['appropriation'] else False or
                query.lower() in str(entry['remarks']).lower() if entry['remarks'] else False or
                query.lower() in str(entry['start_date']).lower() if entry['start_date'] else False or
                query.lower() in str(entry['end_date']).lower() if entry['end_date'] else False or
                query.lower() in str(entry['code']).lower() if entry['code'] else False or
                query.lower() in str(entry['services']).lower() if entry['services'] else False):
                
                
                    matched_entries.append(entry)

    return render(request, 'KwentasApp/projectreports.html', {
        'page_obj': page_obj,
        'all_entries': all_entries,  # Pass all_entries to the template
        'matched_entries': matched_entries,
        'query': query
    })

def download_word(request, project_code):
    # Fetch the specific project entry based on the project code
    _, _, all_entries = get_project_entries()  # Fetch all project entries
    selected_entry = None
    
    for entry in all_entries:
        if entry['code'] == project_code:
            selected_entry = entry
            break

    if not selected_entry:
        return HttpResponse("Project not found", status=404)

    # Create a Word document
    doc = Document()
    doc.add_heading(f"Project Code: {selected_entry['code']}", level=1)
    
    # Add all relevant project details
    doc.add_paragraph(f"Year: {selected_entry['year']}")
    doc.add_paragraph(f"PPA: {selected_entry['ppa']}")
    doc.add_paragraph(f"Implementing Unit: {selected_entry['implementing_unit']}")
    doc.add_paragraph(f"Location: {selected_entry['location']}")
    doc.add_paragraph(f"Appropriation: {selected_entry['appropriation']}")
    doc.add_paragraph(f"Start Date: {selected_entry['start_date']}")
    doc.add_paragraph(f"End Date: {selected_entry['end_date']}")
    doc.add_paragraph(f"Remarks: {selected_entry['remarks']}")
    doc.add_paragraph(f"Remaining Balance: {selected_entry['remaining_total_balance']}")
    doc.add_paragraph(f"Total Spent: {selected_entry['total_disbursements']}")
    doc.add_paragraph(f"Added Budget: {selected_entry['added_budget']}")
    doc.add_paragraph(f"Total Budget: {selected_entry['overall_budget']}")
    doc.add_paragraph(f"Utilization Rate: {selected_entry['utilization_rate']}")

    # Obligations and Budget Data as Lists
    doc.add_heading("Obligation:", level=2)
    for obligation in selected_entry['obligation']:
        doc.add_paragraph(f"{obligation['name']}: {obligation['spent']} ({obligation['date']})", style='List Bullet')

    doc.add_heading("Budget Data:", level=2)
    for budget in selected_entry['budget_data']:
        doc.add_paragraph(f"{budget['name']}: {budget['added_budget']} ({budget['date']})", style='List Bullet')

    # Save the document in a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Return the document as a downloadable file
    response = HttpResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Project_{selected_entry["code"]}.docx'
    
    return response

def get_monthly_expenses():
    result = database.child('Data').get()

    monthly_expenses = {}

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue

            # Check if obligation node exists and aggregate the spent amounts by month
            if 'disbursement' in value:
                for obligation_key, obligation_value in value['disbursement'].items():
                    date = obligation_value.get('date')
                    spent = obligation_value.get('disbursement')

                    if date and spent:
                        month = date[:7]  # Extract YYYY-MM from date
                        spent = float(spent)  # Ensure spent is a number

                        if month in monthly_expenses:
                            monthly_expenses[month] += spent
                        else:
                            monthly_expenses[month] = spent

    return monthly_expenses

def get_monthly_expenses_view(request):
    monthly_expenses = get_monthly_expenses()
    return JsonResponse(monthly_expenses)



def get_daily_expenses_view(request):
    daily_expenses = get_daily_expenses()
    return JsonResponse(daily_expenses)

def get_daily_expenses():
    result = database.child('Data').get()

    daily_expenses = {}

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue

            # Check if disbursement node exists and aggregate the spent amounts by day
            if 'disbursement' in value:
                for disbursement_key, disbursement_value in value['disbursement'].items():
                    date = disbursement_value.get('date')
                    spent = disbursement_value.get('disbursement')

                    if date and spent:
                        day = date[:10]  # Extract YYYY-MM-DD from date
                        spent = float(spent)  # Ensure spent is a float number

                        # Aggregate spent by day
                        if day in daily_expenses:
                            daily_expenses[day] += spent
                        else:
                            daily_expenses[day] = spent

    return daily_expenses

def get_monthly_comparison():
    result = database.child('Data').get()

    monthly_comparison = {
        'months': [],
        'total_disbursements': {},
        'total_obligations': {}
    }

    if result.val():
        for key, value in result.val().items():
            if key == 'placeholder':
                continue

            # Aggregate disbursements
            if 'disbursement' in value:
                for obligation_key, obligation_value in value['disbursement'].items():
                    date = obligation_value.get('date')
                    disbursement_amount = obligation_value.get('disbursement')

                    if date and disbursement_amount:
                        month = date[:7]  # Extract YYYY-MM from date
                        disbursement_amount = float(disbursement_amount)

                        if month in monthly_comparison['total_disbursements']:
                            monthly_comparison['total_disbursements'][month] += disbursement_amount
                        else:
                            monthly_comparison['total_disbursements'][month] = disbursement_amount

            # Aggregate obligations (assumed similar structure, adjust based on actual data)
            if 'obligation' in value:
                for obligation_key, obligation_value in value['obligation'].items():
                    date = obligation_value.get('date')
                    obligation_amount = obligation_value.get('obligation')

                    if date and obligation_amount:
                        month = date[:7]  # Extract YYYY-MM from date
                        obligation_amount = float(obligation_amount)

                        if month in monthly_comparison['total_obligations']:
                            monthly_comparison['total_obligations'][month] += obligation_amount
                        else:
                            monthly_comparison['total_obligations'][month] = obligation_amount

        # Sort months to maintain chronological order
        monthly_comparison['months'] = sorted(set(monthly_comparison['total_disbursements'].keys()) | set(monthly_comparison['total_obligations'].keys()))

    return monthly_comparison

def get_monthly_comparison_view(request):
    monthly_comparison = get_monthly_comparison()
    return JsonResponse(monthly_comparison)

from collections import defaultdict

def get_department_utilization_rate():
    # Get project entries
    entries_below_2024, entries_2024_and_above, all_entries = get_project_entries()

    # Dictionary to store total utilization and count per department
    department_data = defaultdict(lambda: {'total_utilization': 0, 'count': 0})

    # Iterate through all entries
    for entry in all_entries:
        department = entry['implementing_unit']
        utilization_rate = entry.get('utilization_rate', 0)  # Default to 0 if utilization_rate is missing

        # Add utilization rate to the department's total
        department_data[department]['total_utilization'] += float(utilization_rate)
        department_data[department]['count'] += 1

    # Calculate average utilization rate per department
    department_utilization_rates = {}
    for department, data in department_data.items():
        if data['count'] > 0:
            average_utilization = data['total_utilization'] / data['count']
            department_utilization_rates[department] = round(average_utilization, 2)

    return department_utilization_rates

def get_department_utilization_rate_view(request):
    department_utilization_rate = get_department_utilization_rate()
    return JsonResponse(department_utilization_rate)